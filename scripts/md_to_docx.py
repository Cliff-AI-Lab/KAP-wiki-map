"""把 Markdown 文档转 Word .docx（M21 部署清单）。

用 python-docx 简易渲染：标题层级 / 段落 / 表格 / 列表 / 代码块。
不依赖外部 pandoc。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 样式常量 ──
SANS_FONT = '微软雅黑'
MONO_FONT = 'Consolas'
HEADING_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # 深蓝
ACCENT_COLOR = RGBColor(0x3B, 0x82, 0xF6)    # 蓝
MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)


def _set_cell_shading(cell, color_hex: str):
    """给表格单元格加底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(level=min(level, 9))
    run = h.add_run(text)
    run.font.name = SANS_FONT
    run.font.color.rgb = HEADING_COLOR
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    # 中文字体 east-asia
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), SANS_FONT)


def _add_paragraph(doc: Document, text: str, *, mono: bool = False, bullet: bool = False, blockquote: bool = False):
    p = doc.add_paragraph()
    if bullet:
        p.style = 'List Bullet'
    if blockquote:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = MONO_FONT if mono else SANS_FONT
    run.font.size = Pt(10 if mono else 10.5)
    if mono:
        run.font.color.rgb = ACCENT_COLOR
    if blockquote:
        run.font.color.rgb = MUTED_COLOR
        run.font.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), MONO_FONT if mono else SANS_FONT)


def _strip_md_inline(s: str) -> str:
    """去掉行内 markdown 标记（**bold** / `code` / [text](url)）。"""
    s = re.sub(r'\[(.+?)\]\([^)]+\)', r'\1', s)   # link → text
    s = re.sub(r'`([^`]+)`', r'\1', s)             # code
    s = re.sub(r'\*\*([^*]+)\*\*', r'\1', s)       # bold
    s = re.sub(r'\*([^*]+)\*', r'\1', s)           # italic
    s = re.sub(r'^\s*[-*]\s+', '', s)              # leading bullet (handled separately)
    return s.strip()


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$', line))


def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'):   s = s[:-1]
    return [c.strip() for c in s.split('|')]


def _add_table(doc: Document, rows: list[list[str]]):
    """rows[0] 是表头。"""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= len(rows[0]):
                continue
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(_strip_md_inline(cell_text))
            run.font.name = SANS_FONT
            run.font.size = Pt(9.5)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), SANS_FONT)
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_shading(cell, '1F4E79')


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = SANS_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), SANS_FONT)
    rFonts.set(qn('w:ascii'), SANS_FONT)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    in_code = False
    code_buf: list[str] = []
    while i < len(lines):
        line = lines[i]

        # code fence
        if line.startswith('```'):
            if in_code:
                _add_paragraph(doc, '\n'.join(code_buf), mono=True)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line)
            i += 1; continue

        # 表格
        if i + 1 < len(lines) and '|' in line and _is_table_separator(lines[i + 1]):
            rows = [_parse_table_row(line)]
            j = i + 2
            while j < len(lines) and '|' in lines[j] and not _is_table_separator(lines[j]):
                rows.append(_parse_table_row(lines[j]))
                j += 1
            _add_table(doc, rows)
            doc.add_paragraph()  # 表后空行
            i = j; continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            lvl = len(m.group(1))
            _add_heading(doc, _strip_md_inline(m.group(2)), lvl)
            i += 1; continue

        # 引用 blockquote
        if line.startswith('> '):
            _add_paragraph(doc, _strip_md_inline(line[2:]), blockquote=True)
            i += 1; continue

        # 水平线
        if re.match(r'^---+\s*$', line):
            doc.add_paragraph().add_run('—' * 40).font.color.rgb = MUTED_COLOR
            i += 1; continue

        # 列表
        if re.match(r'^[-*]\s+', line):
            _add_paragraph(doc, _strip_md_inline(line), bullet=True)
            i += 1; continue
        if re.match(r'^\d+\.\s+', line):
            num_text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph()
            p.style = 'List Number'
            run = p.add_run(_strip_md_inline(num_text))
            run.font.name = SANS_FONT
            i += 1; continue

        # 空行
        if not line.strip():
            doc.add_paragraph()
            i += 1; continue

        # 普通段落
        _add_paragraph(doc, _strip_md_inline(line))
        i += 1

    doc.save(docx_path)
    print(f'OK: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python md_to_docx.py input.md output.docx')
        sys.exit(1)
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))
