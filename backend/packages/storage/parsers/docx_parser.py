"""DOCX (Word) 文档解析器 — python-docx, 含 mock 回退."""

from __future__ import annotations

import io

from packages.common import get_logger
from packages.storage.parsers.base import BaseParser, ParsedContent

log = get_logger("parser.docx")


class DOCXParser(BaseParser):
    """Word .docx 解析为纯文本 (段落 + 表格)."""

    async def parse(self, raw_bytes: bytes, mime_type: str) -> ParsedContent:
        try:
            return self._parse_real(raw_bytes)
        except Exception as e:
            # V15 修复: 不再静默吞异常 + 假占位入库, 抛出让上游 ingest 显式记 parse_error
            log.warning("docx_parse_failed", error=str(e), mime_type=mime_type)
            raise ValueError(f"DOCX 解析失败: {e}") from e

    def _parse_real(self, raw_bytes: bytes) -> ParsedContent:
        from docx import Document
        d = Document(io.BytesIO(raw_bytes))
        lines: list[str] = []
        # 段落
        for p in d.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)
        # 表格
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    lines.append(line)
        full = "\n".join(lines)
        return ParsedContent(
            text=full,
            metadata={
                "para_count": len(d.paragraphs),
                "table_count": len(d.tables),
                "file_size": len(raw_bytes),
            },
            parser_name="python-docx",
            confidence=0.95,
        )

    def _parse_mock(self, raw_bytes: bytes) -> ParsedContent:
        return ParsedContent(
            text="(DOCX 解析失败 - 文件可能损坏或非标准格式)",
            metadata={"file_size": len(raw_bytes)},
            parser_name="docx_mock",
            confidence=0.0,
        )

    def supported_mime_types(self) -> list[str]:
        # 仅支持 .docx (OOXML); 旧 .doc 是二进制 OLE 格式 python-docx 无法解析,
        # 不注册避免误导用户产出占位文案. 上传 .doc 会走 router.py 的 text_fallback (utf-8) → 乱码,
        # 提示用户用 Word 另存为 .docx.
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]
