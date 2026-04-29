"""知识管理 API 路由。"""

from __future__ import annotations

import structlog
import time
from collections import Counter

log = structlog.get_logger("api.knowledge")
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, HTTPException, Query, Request, UploadFile, File, Form

from api.deps import (
    get_archive_store,
    get_domain_store,
    get_graph_store,
    get_metadata_store,
    get_retriever,
    get_vector_store,
)
from api.middleware.auth import get_current_user
from api.schemas.knowledge import (
    CatalogNode,
    DocumentDetail,
    DocumentSummary,
    KnowledgeStats,
    PaginatedDocuments,
    SearchHit,
    SearchResponse,
)
from packages.common.types import AccessLevel

router = APIRouter(prefix="/knowledge", tags=["知识管理"])


async def _compile_wiki_for_batch(batch_results, project_id: str, domain_store) -> int:
    """公共 Wiki 编译逻辑 — 收集域精炼结果并调用 WikiCompiler。

    返回编译的 Wiki 页数量。供 ingest_demo_data 和 ingest_files 共用。
    """
    from api.deps import get_raw_store, get_wiki_store
    from packages.distillation.wiki_compiler import WikiCompiler

    raw_store = get_raw_store()
    wiki_store = get_wiki_store()
    wiki_compiler = WikiCompiler(raw_store, wiki_store, domain_store)

    domain_refined: dict[str, list[tuple]] = {}
    for r in batch_results:
        if r.decision and r.decision.value == "KEEP" and r.refined_result and r.refined_result.domain_id:
            did = r.refined_result.domain_id
            if did not in domain_refined:
                domain_refined[did] = []
            domain_refined[did].append((r.doc_id, r.refined_result))

    if not domain_refined:
        log.info("wiki_compilation_skipped_no_keep_docs", project_id=project_id)
        return 0

    log.info("wiki_compilation_starting", project_id=project_id, domains=len(domain_refined))
    wiki_pages = await wiki_compiler.compile_project(project_id, domain_results=domain_refined)
    log.info("wiki_compilation_done", pages=len(wiki_pages), domains=len(domain_refined))
    return len(wiki_pages)


# RBAC 权限等级排序
_ACCESS_LEVEL_RANK: dict[str, int] = {
    AccessLevel.PUBLIC.value: 0,
    AccessLevel.INTERNAL.value: 1,
    AccessLevel.CONFIDENTIAL.value: 2,
    AccessLevel.SECRET.value: 3,
}


# ── 辅助函数 ──────────────────────────────────────────


def _rbac_filter(docs: list[dict], user_access_level: str, user_department: str | None) -> list[dict]:
    """按用户权限等级过滤文档列表。"""
    user_rank = _ACCESS_LEVEL_RANK.get(user_access_level, 1)
    filtered = []
    for d in docs:
        doc_access = d.get("access_level", AccessLevel.INTERNAL.value)
        doc_rank = _ACCESS_LEVEL_RANK.get(doc_access, 1)
        if user_rank < doc_rank:
            continue
        # CONFIDENTIAL+ 文档需部门匹配
        if doc_rank >= _ACCESS_LEVEL_RANK[AccessLevel.CONFIDENTIAL.value]:
            doc_dept = d.get("department_id", "")
            if doc_dept and user_department and doc_dept != user_department:
                continue
        filtered.append(d)
    return filtered


def _str_field(doc: dict, key: str, default: str = "") -> str:
    """安全获取字段的字符串表示。"""
    val = doc.get(key)
    if val is None:
        return default
    return str(val)


def _parse_keywords(raw) -> list[str]:
    """将 keywords 字段解析为列表。"""
    if isinstance(raw, list):
        return raw
    if isinstance(raw, str) and raw:
        return [k.strip() for k in raw.split(",") if k.strip()]
    return []


def _build_catalog_tree(
    tree: dict, parent_path: str = ""
) -> list[CatalogNode]:
    """递归构建目录树节点列表。"""
    nodes: list[CatalogNode] = []
    for name, info in sorted(tree.items()):
        path = f"{parent_path}/{name}" if parent_path else name
        children = _build_catalog_tree(info["_children"], path)
        nodes.append(
            CatalogNode(
                path=path,
                name=name,
                doc_count=info["_count"],
                children=children,
            )
        )
    return nodes


# ── 端点 ──────────────────────────────────────────────


@router.get("/stats", response_model=KnowledgeStats)
async def get_stats(
    request: Request,
    project_id: str = Query(default="default", description="项目ID"),
) -> KnowledgeStats:
    """获取知识库统计信息。"""
    meta = get_metadata_store()
    vec = get_vector_store()
    graph = get_graph_store()
    domain_store = get_domain_store()
    user = get_current_user(request)

    all_docs = await meta.list_documents(org_id=project_id)

    decision_counts = Counter(d.get("decision", "") for d in all_docs)
    status_counts = Counter(d.get("status", "") for d in all_docs)
    type_counts = Counter(
        _str_field(d, "doc_type", "其他") for d in all_docs
    )
    source_counts = Counter(
        _str_field(d, "source_system", "unknown") for d in all_docs
    )

    domains = domain_store.list_domains(project_id=project_id)
    return KnowledgeStats(
        total_documents=len(all_docs),
        kept=decision_counts.get("KEEP", 0),
        archived=decision_counts.get("ARCHIVE", 0),
        discarded=decision_counts.get("DISCARD", 0),
        pending_review=status_counts.get("PENDING_REVIEW", 0),
        vector_chunks=vec.chunk_count,
        knowledge_domains=len(domains),
        doc_cards=domain_store.card_count,
        graph_nodes=graph.node_count,
        graph_edges=graph.edge_count,
        by_doc_type=dict(type_counts),
        by_source_system=dict(source_counts),
    )


@router.get("/catalog", response_model=list[CatalogNode])
async def get_catalog(
    request: Request,
    project_id: str = Query(default="default", description="项目ID"),
) -> list[CatalogNode]:
    """获取知识目录树。"""
    meta = get_metadata_store()
    user = get_current_user(request)
    all_docs = await meta.list_documents(org_id=project_id)
    all_docs = _rbac_filter(all_docs, user.access_level, user.department_id)

    tree: dict = {}
    for d in all_docs:
        path = _str_field(d, "category_path") or "未分类"
        parts = [p for p in path.split("/") if p]
        current = tree
        for part in parts:
            if part not in current:
                current[part] = {"_count": 0, "_children": {}}
            current[part]["_count"] += 1
            current = current[part]["_children"]

    return _build_catalog_tree(tree)


@router.get("/documents", response_model=PaginatedDocuments)
async def list_documents(
    request: Request,
    status: Optional[str] = None,
    decision: Optional[str] = None,
    doc_type: Optional[str] = None,
    source_system: Optional[str] = None,
    category_path: Optional[str] = None,
    project_id: str = Query(default="default", description="项目ID"),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
) -> PaginatedDocuments:
    """列出文档（支持过滤与分页）。"""
    meta = get_metadata_store()
    user = get_current_user(request)
    docs = await meta.list_documents(status=status, decision=decision, org_id=project_id)

    # RBAC 权限过滤
    docs = _rbac_filter(docs, user.access_level, user.department_id)

    # 额外过滤
    if doc_type:
        docs = [d for d in docs if _str_field(d, "doc_type") == doc_type]
    if source_system:
        docs = [
            d for d in docs if _str_field(d, "source_system") == source_system
        ]
    if category_path:
        docs = [
            d
            for d in docs
            if _str_field(d, "category_path").startswith(category_path)
        ]

    # 按 updated_at 降序排序
    docs.sort(key=lambda d: d.get("updated_at") or "", reverse=True)

    total = len(docs)
    start = (page - 1) * page_size
    end = start + page_size
    page_docs = docs[start:end]

    summaries = [
        DocumentSummary(
            id=_str_field(d, "id"),
            title=_str_field(d, "title"),
            doc_type=_str_field(d, "doc_type"),
            decision=_str_field(d, "decision"),
            status=_str_field(d, "status"),
            kpi_retain=d.get("kpi_retain"),
            source_system=_str_field(d, "source_system"),
            summary=(_str_field(d, "summary"))[:200],
            category_path=_str_field(d, "category_path"),
            keywords=_parse_keywords(d.get("keywords")),
            created_at=_str_field(d, "created_at"),
            updated_at=_str_field(d, "updated_at"),
        )
        for d in page_docs
    ]

    return PaginatedDocuments.build(summaries, total, page, page_size)


@router.get("/documents/{doc_id}", response_model=DocumentDetail)
async def get_document(request: Request, doc_id: str) -> DocumentDetail:
    """获取文档详情。"""
    meta = get_metadata_store()
    graph = get_graph_store()
    user = get_current_user(request)

    doc = await meta.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # RBAC 权限检查
    if not _rbac_filter([doc], user.access_level, user.department_id):
        raise HTTPException(status_code=403, detail="无权访问该文档")

    # 从图谱获取实体
    entities: list[str] = []
    related_doc_ids: list[str] = []
    try:
        entity_names = await graph.get_doc_entities(doc_id)
        entities = entity_names or []
        # 通过实体查找关联文档
        seen: set[str] = set()
        for ent in entities[:10]:  # 限制查询数量
            related = await graph.find_related_docs(ent, max_hops=1)
            for rid in related:
                if rid != doc_id and rid not in seen:
                    seen.add(rid)
                    related_doc_ids.append(rid)
    except Exception:
        pass  # 图谱查询失败不影响主流程

    return DocumentDetail(
        id=_str_field(doc, "id"),
        title=_str_field(doc, "title"),
        doc_type=_str_field(doc, "doc_type"),
        decision=_str_field(doc, "decision"),
        status=_str_field(doc, "status"),
        kpi_retain=doc.get("kpi_retain"),
        source_system=_str_field(doc, "source_system"),
        summary=_str_field(doc, "summary"),
        keywords=_parse_keywords(doc.get("keywords")),
        judge_reasoning=doc.get("judge_reasoning"),
        department_id=_str_field(doc, "department_id"),
        access_level=_str_field(doc, "access_level", "INTERNAL"),
        category_path=_str_field(doc, "category_path"),
        created_at=_str_field(doc, "created_at"),
        updated_at=_str_field(doc, "updated_at"),
        ingested_at=_str_field(doc, "ingested_at"),
        entities=entities,
        related_doc_ids=related_doc_ids,
    )


@router.get("/search", response_model=SearchResponse)
async def search_knowledge(
    request: Request,
    q: str = Query(..., min_length=1, max_length=500, description="搜索关键词"),
    top_k: int = Query(default=10, ge=1, le=50),
    category_path: Optional[str] = None,
    project_id: str = Query(default="default", description="项目ID"),
) -> SearchResponse:
    """知识搜索（语义检索）。"""
    start = time.time()
    retriever = get_retriever()
    user = get_current_user(request)
    results = await retriever.search(
        query=q, top_k=top_k, target_category=category_path,
        org_id=project_id, user_access_level=user.access_level,
        user_department=user.department_id,
    )
    latency = int((time.time() - start) * 1000)

    return SearchResponse(
        query=q,
        total_hits=len(results),
        results=[
            SearchHit(
                doc_id=r.doc_id,
                chunk_id=r.chunk_id,
                title=r.title,
                content=r.content[:300],
                score=r.score,
                source_system=r.source_system,
                category_path=r.category_path,
            )
            for r in results
        ],
        latency_ms=latency,
    )


# ── 知识域端点（Skills 模式） ────────────────────────


@router.get("/domains")
async def list_domains(
    project_id: str = Query(default="default", description="项目ID"),
) -> dict:
    """查看知识域目录（Skills 模式核心 — LLM 路由时读的目录）。"""
    ds = get_domain_store()
    domains = ds.list_domains(project_id=project_id)
    catalog_text = ds.get_domain_catalog_text(project_id=project_id)
    return {
        "total_domains": len(domains),
        "total_doc_cards": ds.card_count,
        "domains": [
            {
                "domain_id": d.domain_id,
                "name": d.name,
                "parent_id": d.parent_id,
                "description": d.description,
                "doc_count": d.doc_count,
                "is_system": d.is_system,
            }
            for d in domains
        ],
        "catalog_text_for_llm": catalog_text,
    }


@router.post("/reload-domains")
async def reload_domains() -> dict:
    """从 PG 重新加载文档卡数据到内存。"""
    ds = get_domain_store()
    count = await ds.reload_from_pg()
    return {"reloaded": count}


@router.get("/debug-cards")
async def debug_cards() -> dict:
    """临时调试：查看内存中的文档卡 domain_id 分布。"""
    ds = get_domain_store()
    domain_counts: dict[str, int] = {}
    for doc_id, card in ds._doc_cards.items():
        pid = ds._card_project.get(doc_id, "?")
        key = f"{pid}:{card.domain_id}"
        domain_counts[key] = domain_counts.get(key, 0) + 1
    return {"total_cards": len(ds._doc_cards), "distribution": domain_counts}


@router.get("/graph-overview")
async def get_graph_overview(
    project_id: str = Query(default="default", description="项目ID"),
    domain_id: str = Query(default="", description="知识体系分支 ID（按分支过滤图谱）"),
) -> dict:
    """V8: 图谱概览 — 返回 Entity 节点(含doc_ids) + RELATES 连线(含weight)。

    按知识体系分支逐层展开，不加载全图。
    """
    graph = get_graph_store()
    nodes, edges = await graph.get_entities_by_domain(domain_id)

    return {
        "node_count": len(nodes),
        "edge_count": len(edges),
        "nodes": nodes[:200],
        "edges": edges[:500],
        "domain_id": domain_id,
    }


@router.get("/graph-doc-view")
async def get_graph_doc_view(
    project_id: str = Query(default="default", description="项目ID"),
    domain_id: str = Query(default="", description="知识体系分支 ID"),
) -> dict:
    """V8 新增：文档视角图谱 — 节点=文档，边=共享实体数。"""
    graph = get_graph_store()
    nodes, edges = await graph.get_doc_view_graph(domain_id)

    # 用文档标题替换 doc_id 作为 label
    meta = get_metadata_store()
    for node in nodes:
        doc_meta = await meta.get_document(node["id"])
        if doc_meta:
            node["label"] = doc_meta.get("title", node["id"])

    return {
        "node_count": len(nodes),
        "edge_count": len(edges),
        "nodes": nodes,
        "edges": edges,
        "domain_id": domain_id,
    }


# ── 影子归档端点 ──────────────────────────────────────


@router.get("/archive")
async def list_archived_documents() -> list[dict]:
    """列出所有影子归档文档。"""
    archive = get_archive_store()
    return await archive.list_archived()


@router.post("/archive/{doc_id}/restore")
async def restore_document(doc_id: str) -> dict:
    """从影子归档恢复文档。"""
    archive = get_archive_store()
    meta = get_metadata_store()

    data = await archive.restore_document(doc_id)
    if not data:
        raise HTTPException(status_code=404, detail="归档文档不存在或已过期")

    # 恢复元数据状态为 ACTIVE / KEEP
    doc_record = data.get("metadata", {})
    doc_record["id"] = doc_id
    doc_record["status"] = "ACTIVE"
    doc_record["decision"] = "KEEP"
    await meta.upsert_document(doc_record)

    # 清理归档记录
    await archive.delete_archive(doc_id)

    return {"doc_id": doc_id, "status": "restored"}


# ── 人工审核队列端点 ──────────────────────────────────


@router.get("/review-queue")
async def list_review_queue(
    status: str = "PENDING",
    project_id: str = Query(default="default", description="项目ID"),
) -> list[dict]:
    """列出人工审核队列（按项目过滤）。"""
    meta = get_metadata_store()
    return await meta.list_review_queue(status=status, project_id=project_id)


@router.post("/review-queue/{doc_id}/resolve")
async def resolve_review(
    doc_id: str,
    final_decision: str = Query(..., description="最终决策: KEEP/ARCHIVE/DISCARD"),
    reviewer: str = Query(default="admin", description="审核人"),
) -> dict:
    """审核员处理审核队列条目。"""
    if final_decision not in ("KEEP", "ARCHIVE", "DISCARD"):
        raise HTTPException(status_code=400, detail="无效决策，必须为 KEEP/ARCHIVE/DISCARD")

    meta = get_metadata_store()
    success = await meta.resolve_review(doc_id, final_decision, reviewer)
    if not success:
        raise HTTPException(status_code=404, detail="审核条目不存在或已处理")

    return {"doc_id": doc_id, "final_decision": final_decision, "reviewer": reviewer}


# ── Demo 数据灌入 ──────────────────────────────────────


@router.post("/ingest-demo")
async def ingest_demo_data(
    request: Request,
    project_id: str = Query(default="default", description="项目ID"),
    force: bool = Query(default=False, description="强制重新灌入（清除旧数据）"),
    industry: str = Query(default="auto", description="行业类型: auto(自动检测)/energy/default"),
) -> dict:
    """一键灌入多源 Mock 数据（飞书+钉钉+企微 → 蒸馏 → 入库）。"""
    from packages.connectors.feishu import FeishuConnector
    from packages.connectors.dingtalk import DingTalkConnector
    from packages.connectors.wecom import WeComConnector
    from packages.distillation.pipeline import arun_pipeline
    from packages.common.types import Decision, DocumentCard
    from packages.storage.chunker import chunk_document
    from packages.storage.embedder import aembed_texts

    meta = get_metadata_store()
    vec = get_vector_store()
    graph = get_graph_store()
    domain_store = get_domain_store()
    archive = get_archive_store()

    # V6: 生成项目级知识域列表供 Refiner 使用
    project_domain_list = domain_store.get_refiner_domain_list(project_id=project_id)

    # 检查是否已有数据
    existing = await meta.list_documents(org_id=project_id)
    if existing and not force:
        return {"status": "skip", "message": f"知识库已有 {len(existing)} 篇文档，跳过灌入。传 ?force=true 强制重灌"}

    # 强制重灌：清除所有旧数据（含 V11 新增的 raw/wiki store）
    if existing and force:
        await meta.clear_all()
        await vec.clear_all()
        await graph.clear_all()
        await domain_store.clear_doc_cards()
        # 清除 RawStore 和 WikiStore
        from api.deps import get_raw_store, get_wiki_store
        try:
            raw_store = get_raw_store()
            await raw_store.clear_all(project_id=project_id)
        except Exception:
            pass
        try:
            wiki_store = get_wiki_store()
            await wiki_store.clear_all(project_id=project_id)
        except Exception:
            pass
        # 清除本地归档
        import shutil
        from pathlib import Path
        archive_dir = Path("data/archive")
        if archive_dir.exists():
            shutil.rmtree(archive_dir)
            archive_dir.mkdir(parents=True, exist_ok=True)

    # Phase 1: 多源采集
    # V9: 能源行业项目使用专用能源 mock 数据
    # V9: 判断是否使用能源行业 mock 数据
    if industry == 'auto':
        from api.deps import get_project_store
        project_store = get_project_store()
        project_info = project_store.get_project(project_id)
        use_energy_data = bool(project_info and project_info.get('industry_code') == 'energy')
    else:
        use_energy_data = (industry == 'energy')

    all_docs = []
    source_counts = {}

    if use_energy_data:
        # 能源行业：使用30篇能源行业专用文档
        from packages.connectors.energy_mock_data import (
            get_feishu_energy_docs, get_dingtalk_energy_docs, get_wecom_energy_docs,
        )
        from packages.common.types import RawDocument, SourceSystem
        for name, getter, src_sys in [
            ("feishu", get_feishu_energy_docs, SourceSystem.FEISHU),
            ("dingtalk", get_dingtalk_energy_docs, SourceSystem.DINGTALK),
            ("wecom", get_wecom_energy_docs, SourceSystem.WECOM),
        ]:
            raw_dicts = getter()
            docs = []
            for d in raw_dicts:
                doc = RawDocument(
                    doc_id=d["doc_id"], title=d["title"],
                    content=d["content"], source_system=src_sys,
                    source_id=d.get("source_id", ""), org_id=project_id,
                    created_at=d.get("created_at"), updated_at=d.get("updated_at"),
                    created_by=d.get("created_by", ""), last_modifier=d.get("last_modifier", ""),
                )
                docs.append(doc)
            all_docs.extend(docs)
            source_counts[name] = len(docs)
        print(f"[energy_mock] loaded {len(all_docs)} energy docs")
    else:
        # 默认：使用原有 AI 质检 mock 数据
        for name, ConnClass in [("feishu", FeishuConnector), ("dingtalk", DingTalkConnector), ("wecom", WeComConnector)]:
            conn = ConnClass()
            await conn.connect()
            docs = []
            async for doc in conn.fetch_documents():
                docs.append(doc)
            all_docs.extend(docs)
            source_counts[name] = len(docs)

    # V6: 强制设置所有文档的 org_id 为当前项目
    for doc in all_docs:
        doc.org_id = project_id

    # V11: Phase 1.5 — 保存原始文档到 RawStore（不可变层，**保存原文**）
    from api.deps import get_raw_store, get_wiki_store
    raw_store = get_raw_store()
    for doc in all_docs:
        try:
            await raw_store.save_raw(doc, project_id=project_id)
        except Exception as e:
            log.warning("raw_store_save_failed", doc_id=doc.doc_id, error=str(e))
    log.info("raw_store_saved", count=len(all_docs))

    # M2 #2: W1 脱敏 hook（决策书 §5.4 D10 工位嵌入）
    # 在 RawStore 保存原文后，对 doc.content 就地脱敏 + 映射持久化；
    # pipeline 后续基于脱敏文做嵌入 + 入库（vec_redacted 默认）
    try:
        from packages.sensitive.ingest_hook import redact_and_persist_doc
        from packages.sensitive.mapping_store import get_mapping_store
        mstore = get_mapping_store()
        await mstore.initialize()
        redact_count = 0
        for doc in all_docs:
            r = await redact_and_persist_doc(doc, mapping_store=mstore)
            if r.tokens:
                redact_count += 1
        log.info("ingest_redaction_done", redacted_docs=redact_count, total=len(all_docs))
    except Exception as e:
        # 脱敏失败不阻断 ingest（轻量化兜底）；生产应升级为告警
        log.warning("ingest_redaction_failed_continue", error=str(e))

    # Phase 2: 蒸馏（传入项目级知识域列表，**基于脱敏文**）
    batch = await arun_pipeline(all_docs, domain_list_text=project_domain_list)
    doc_map = {d.doc_id: d for d in all_docs}

    # Phase 3: 入库（处理所有决策类型：KEEP/ARCHIVE/DISCARD）
    total_chunks = 0
    kept_docs = 0
    archived_docs = 0
    discarded_docs = 0
    review_enqueued = 0
    all_bm25_entries: list[dict] = []

    # OPT-01: 使用自动目录生成器替代简单拼接
    from packages.distillation.auto_cataloger import get_auto_cataloger
    cataloger = get_auto_cataloger()

    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        if not doc or not r.decision:
            continue

        cat = cataloger.generate_category_path(
            key_topics=r.librarian_result.key_topics if r.librarian_result else [],
            domain_id=r.refined_result.domain_id if r.refined_result else "",
            doc_type=r.librarian_result.doc_type.value if r.librarian_result else "",
        )
        summary = r.refined_result.summary if r.refined_result else ""
        kw = ",".join(r.refined_result.keywords) if r.refined_result else ""
        domain_id = r.refined_result.domain_id if r.refined_result else ""
        doc_type_str = r.librarian_result.doc_type.value if r.librarian_result else "其他"

        # 确定文档状态
        if r.needs_review:
            doc_status = "PENDING_REVIEW"
        elif r.decision == Decision.KEEP:
            doc_status = "ACTIVE"
        elif r.decision == Decision.ARCHIVE:
            doc_status = "ARCHIVED"
        else:
            doc_status = "DISCARDED"

        # M1 ISS DataScope 激活：从 request user context 拉 dept_id/created_by
        # 系统导入（demo / mock 连接器）时 user 是 anonymous，dept_id=None / created_by=""
        # 让 retriever DataScope 过滤端透明放行（M0 兼容）；真实用户登录时填实
        user = getattr(request.state, "user", None)
        ingest_dept_id = getattr(user, "dept_id", None) if user else None
        ingest_created_by = getattr(user, "user_id", "") if user else ""
        if ingest_created_by == "anonymous":
            ingest_created_by = ""

        # 所有文档都写入元数据（不再跳过 ARCHIVE/DISCARD）
        await meta.upsert_document({
            "id": doc.doc_id,
            "title": doc.title,
            "source_system": doc.source_system.value,
            "doc_type": doc_type_str,
            "version_id": None,
            "status": doc_status,
            "decision": r.decision.value,
            "kpi_retain": r.judge_result.kpi_retain if r.judge_result else None,
            "summary": summary,
            "keywords": kw,
            "category_path": cat,
            "org_id": doc.org_id,
            "created_at": None,
            "updated_at": None,
            "access_level": doc.metadata.get("access_level", "INTERNAL"),
            "department_id": doc.metadata.get("department_id", ""),
            "dept_id": ingest_dept_id,             # M1 ISS DataScope int
            "created_by": ingest_created_by,        # M1 ISS DataScope SELF
            "judge_reasoning": r.judge_result.reasoning.model_dump() if r.judge_result else None,
        })

        # BUG-004 修复：低置信度文档加入审核队列
        if r.needs_review and r.judge_result:
            await meta.enqueue_review(
                doc_id=doc.doc_id,
                proposed_decision=r.decision.value,
                confidence=r.judge_result.confidence,
                kpi_retain=r.judge_result.kpi_retain,
                reason=r.judge_result.summary or "低置信度自动进入审核",
            )
            # M1 W4 写入侧：双写到 4×6 矩阵审核台（决策书 §5.2 W4 必审）
            from api.deps import get_governance_queue_store
            from packages.governance.distillation_hook import enqueue_low_confidence_review
            try:
                await enqueue_low_confidence_review(
                    store=get_governance_queue_store(),
                    project_id=project_id,
                    doc_id=doc.doc_id,
                    doc_title=doc.title,
                    confidence=r.judge_result.confidence,
                    proposed_decision=r.decision.value,
                    reason=r.judge_result.summary or "低置信度自动进入审核",
                )
            except Exception as e:
                # 双写失败不影响 V15 主路径，仅记日志
                log.warning("matrix_review_enqueue_failed", doc_id=doc.doc_id, error=str(e))
            review_enqueued += 1

        # BUG-003 修复：DISCARD 文档写入影子归档（30天可恢复）
        if r.decision == Decision.DISCARD and not r.is_noise:
            await archive.archive_document(
                doc_id=doc.doc_id,
                content=doc.content,
                metadata={
                    "title": doc.title,
                    "source_system": doc.source_system.value,
                    "doc_type": doc_type_str,
                    "kpi_retain": r.judge_result.kpi_retain if r.judge_result else None,
                    "summary": summary,
                    "org_id": doc.org_id,
                },
            )
            discarded_docs += 1
            continue  # DISCARD 不做切片和图谱

        if r.decision == Decision.ARCHIVE:
            archived_docs += 1
            continue  # ARCHIVE 文档只保留元数据，不做切片

        # KEEP 文档：切片 + 向量 + 图谱 + 文档卡
        kept_docs += 1
        chunks = chunk_document(
            doc_id=doc.doc_id,
            content=doc.content,
            category_path=cat,
            doc_type=doc_type_str,
            source_system=doc.source_system.value,
            updated_at=doc.updated_at,
            org_id=doc.org_id,
            domain_id=domain_id,
        )

        if chunks:
            for c in chunks:
                c.domain_id = domain_id

            texts = [c.content for c in chunks]
            embeddings = await aembed_texts(texts)
            for chunk, emb in zip(chunks, embeddings):
                chunk.embedding = emb
            await vec.insert_chunks(chunks)
            total_chunks += len(chunks)
            for c in chunks:
                all_bm25_entries.append({"chunk_id": c.chunk_id, "doc_id": c.doc_id, "content": c.content})

        if r.refined_result:
            # V7: 图谱轻量化 — 只写 Entity+RELATES，不写 Document/Section 节点
            await graph.add_entities_and_relations(
                doc.doc_id, r.refined_result.entities, r.refined_result.relations,
                domain_id=domain_id,
            )

            doc_card = DocumentCard(
                doc_id=doc.doc_id,
                title=doc.title,
                domain_id=domain_id,
                description=r.refined_result.doc_description or summary[:200],
                key_elements=r.refined_result.key_elements,
                keywords=r.refined_result.keywords,
            )
            await domain_store.upsert_doc_card(doc_card, project_id=project_id)

    # 刷新图谱计数
    await graph.refresh_counts()

    # V11: Phase 4 — Wiki 编译（将域内文档编译为知识 Wiki 页）
    try:
        await _compile_wiki_for_batch(batch.results, project_id, domain_store)
    except Exception as e:
        log.warning("wiki_compilation_failed", error=str(e))

    # 构建 BM25 索引
    from api.deps import get_keyword_scorer
    bm25 = get_keyword_scorer()
    bm25.build_index(all_bm25_entries)

    # OPT-13: 清除结果缓存（数据已变更）
    from api.deps import get_result_cache
    try:
        get_result_cache().invalidate_all()
    except AssertionError:
        pass  # 缓存未初始化时忽略

    # 灌入完成后从 PG 同步文档卡到内存（确保域计数准确）
    try:
        await domain_store.reload_from_pg()
    except Exception:
        pass

    # 构建每个文档的详细结果（供前端去噪审核展示）
    doc_map = {d.doc_id: d for d in all_docs}
    doc_results = []
    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        doc_results.append({
            "doc_id": r.doc_id,
            "title": doc.title if doc else r.doc_id,
            "decision": r.decision.value if r.decision else "UNKNOWN",
            "domain_id": r.refined_result.domain_id if r.refined_result else "",
            "summary": (r.refined_result.summary if r.refined_result else "")[:200],
            "doc_type": r.librarian_result.doc_type.value if r.librarian_result else "未知",
            "entity_count": len(r.refined_result.entities) if r.refined_result else 0,
            "keyword_count": len(r.refined_result.keywords) if r.refined_result else 0,
            "confidence": r.judge_result.confidence if r.judge_result else 0,
            "needs_review": r.needs_review,
        })

    return {
        "status": "ok",
        "total_collected": len(all_docs),
        "source_counts": source_counts,
        "distillation": {
            "kept": batch.kept,
            "archived": batch.archived,
            "discarded": batch.discarded,
            "noise_filtered": batch.noise_filtered,
        },
        "storage": {
            "documents": kept_docs,
            "archived_documents": archived_docs,
            "discarded_archived": discarded_docs,
            "review_enqueued": review_enqueued,
            "vector_chunks": total_chunks,
            "knowledge_domains": domain_store.domain_count,
            "doc_cards": domain_store.card_count,
            "graph_nodes": graph.node_count,
            "graph_edges": graph.edge_count,
        },
        "documents": doc_results,
    }


@router.post("/ingest")
async def ingest_files(
    request: Request,
    files: list[UploadFile] = File(..., description="上传的文件列表"),
    project_id: str = Form(default="default", description="项目ID"),
) -> dict:
    """上传真实文件 → 解析 → 蒸馏 → 入库。

    支持: .txt, .md, .docx, .pdf
    """
    import hashlib
    from datetime import datetime, timezone
    from packages.distillation.pipeline import arun_pipeline
    from packages.common.types import Decision, DocumentCard, RawDocument, SourceSystem
    from packages.storage.chunker import chunk_document
    from packages.storage.embedder import aembed_texts

    meta = get_metadata_store()
    vec = get_vector_store()
    graph = get_graph_store()
    domain_store = get_domain_store()
    project_domain_list = domain_store.get_refiner_domain_list(project_id=project_id)

    # 解析上传的文件为 RawDocument
    raw_docs: list[RawDocument] = []
    parse_errors: list[str] = []
    for f in files:
        try:
            content = await _parse_upload_file(f)
            if not content or len(content.strip()) < 10:
                parse_errors.append(f"{f.filename}: 内容为空或过短")
                continue
            content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
            name_hash = hashlib.md5((f.filename or '').encode()).hexdigest()[:8]
            doc_id = f"upload_{name_hash}_{content_hash}"
            raw_docs.append(RawDocument(
                doc_id=doc_id,
                title=Path(f.filename or "untitled").stem,
                content=content,
                source_system=SourceSystem.LOCAL if hasattr(SourceSystem, 'LOCAL') else SourceSystem.FEISHU,
                org_id=project_id,
                created_at=datetime.now(timezone.utc),
                updated_at=datetime.now(timezone.utc),
                file_size=len(content),
                metadata={"original_filename": f.filename},
            ))
        except Exception as e:
            parse_errors.append(f"{f.filename}: {str(e)}")

    if not raw_docs:
        raise HTTPException(status_code=400, detail=f"无可处理的文件。{'; '.join(parse_errors)}")

    # 蒸馏
    batch = await arun_pipeline(raw_docs, domain_list_text=project_domain_list)
    doc_map = {d.doc_id: d for d in raw_docs}

    # 入库（复用 ingest-demo 的逻辑）
    from packages.distillation.auto_cataloger import get_auto_cataloger
    cataloger = get_auto_cataloger()
    total_chunks = 0
    kept_docs = 0
    all_bm25_entries: list[dict] = []

    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        if not doc or not r.decision:
            continue

        cat = cataloger.generate_category_path(
            key_topics=r.librarian_result.key_topics if r.librarian_result else [],
            domain_id=r.refined_result.domain_id if r.refined_result else "",
            doc_type=r.librarian_result.doc_type.value if r.librarian_result else "",
        )
        summary = r.refined_result.summary if r.refined_result else ""
        kw = ",".join(r.refined_result.keywords) if r.refined_result else ""
        domain_id = r.refined_result.domain_id if r.refined_result else ""
        doc_type_str = r.librarian_result.doc_type.value if r.librarian_result else "其他"

        doc_status = "ACTIVE" if r.decision == Decision.KEEP else (
            "PENDING_REVIEW" if r.needs_review else (
                "ARCHIVED" if r.decision == Decision.ARCHIVE else "DISCARDED"
            )
        )

        await meta.upsert_document({
            "id": doc.doc_id, "title": doc.title,
            "source_system": "local", "doc_type": doc_type_str,
            "version_id": None, "status": doc_status,
            "decision": r.decision.value,
            "kpi_retain": r.judge_result.kpi_retain if r.judge_result else None,
            "summary": summary, "keywords": kw,
            "category_path": cat, "org_id": doc.org_id,
            "created_at": None, "updated_at": None,
            "access_level": "INTERNAL", "department_id": "",
            # M1 ISS DataScope 激活：上传文件时关联当前登录用户
            "dept_id": getattr(getattr(request.state, "user", None), "dept_id", None),
            "created_by": (
                getattr(getattr(request.state, "user", None), "user_id", "") or ""
            ).replace("anonymous", ""),
            "judge_reasoning": r.judge_result.reasoning.model_dump() if r.judge_result else None,
        })

        if r.needs_review and r.judge_result:
            await meta.enqueue_review(
                doc_id=doc.doc_id, proposed_decision=r.decision.value,
                confidence=r.judge_result.confidence,
                kpi_retain=r.judge_result.kpi_retain,
                reason=r.judge_result.summary or "低置信度",
            )
            # M1 W4 写入侧：双写到 4×6 矩阵审核台
            from api.deps import get_governance_queue_store
            from packages.governance.distillation_hook import enqueue_low_confidence_review
            try:
                await enqueue_low_confidence_review(
                    store=get_governance_queue_store(),
                    project_id=project_id,
                    doc_id=doc.doc_id,
                    doc_title=doc.title,
                    confidence=r.judge_result.confidence,
                    proposed_decision=r.decision.value,
                    reason=r.judge_result.summary or "低置信度",
                )
            except Exception as e:
                log.warning("matrix_review_enqueue_failed", doc_id=doc.doc_id, error=str(e))

        if r.decision != Decision.KEEP:
            continue

        kept_docs += 1
        chunks = chunk_document(
            doc_id=doc.doc_id, content=doc.content,
            category_path=cat, doc_type=doc_type_str,
            source_system="local", updated_at=doc.updated_at,
            org_id=doc.org_id,
            domain_id=domain_id,
        )
        if chunks:
            embeddings = await aembed_texts([c.content for c in chunks])
            for chunk, emb in zip(chunks, embeddings):
                chunk.embedding = emb
            await vec.insert_chunks(chunks)
            total_chunks += len(chunks)
            for c in chunks:
                all_bm25_entries.append({"chunk_id": c.chunk_id, "doc_id": c.doc_id, "content": c.content})

        if r.refined_result:
            # V7: 图谱轻量化 — 只写 Entity+RELATES
            await graph.add_entities_and_relations(
                doc.doc_id, r.refined_result.entities, r.refined_result.relations,
                domain_id=domain_id,
            )
            doc_card = DocumentCard(
                doc_id=doc.doc_id, title=doc.title, domain_id=domain_id,
                description=r.refined_result.doc_description or summary[:200],
                key_elements=r.refined_result.key_elements,
                keywords=r.refined_result.keywords,
            )
            await domain_store.upsert_doc_card(doc_card, project_id=project_id)

    await graph.refresh_counts()

    # 增量更新 BM25
    if all_bm25_entries:
        from api.deps import get_keyword_scorer
        get_keyword_scorer().add_chunks(all_bm25_entries)

    # V11.2: 保存原始文档到 RawStore（不可变层），确保 Wiki 编译时可溯源标题
    from api.deps import get_raw_store
    raw_store = get_raw_store()
    for doc in raw_docs:
        try:
            await raw_store.save_raw(doc, project_id=project_id)
        except Exception as e:
            log.warning("raw_store_save_failed", doc_id=doc.doc_id, error=str(e))

    # V11.2: Wiki 编译（per-source + domain + index）
    wiki_pages_compiled = 0
    try:
        wiki_pages_compiled = await _compile_wiki_for_batch(batch.results, project_id, domain_store)
    except Exception as e:
        log.warning("ingest_wiki_compilation_failed", error=str(e), exc_info=True)

    # 构建每个文档的详细结果（前端可直接展示）
    doc_results = []
    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        doc_results.append({
            "doc_id": r.doc_id,
            "title": doc.title if doc else r.doc_id,
            "decision": r.decision.value if r.decision else "UNKNOWN",
            "domain_id": r.refined_result.domain_id if r.refined_result else "",
            "summary": (r.refined_result.summary if r.refined_result else "")[:200],
            "doc_type": r.librarian_result.doc_type.value if r.librarian_result else "未知",
            "entity_count": len(r.refined_result.entities) if r.refined_result else 0,
            "keyword_count": len(r.refined_result.keywords) if r.refined_result else 0,
            "confidence": r.judge_result.confidence if r.judge_result else 0,
            "needs_review": r.needs_review,
        })

    return {
        "status": "ok",
        "total_uploaded": len(files),
        "parsed": len(raw_docs),
        "parse_errors": parse_errors,
        "distillation": {
            "kept": batch.kept, "archived": batch.archived,
            "discarded": batch.discarded, "noise_filtered": batch.noise_filtered,
        },
        "storage": {
            "documents_kept": kept_docs, "vector_chunks": total_chunks,
            "wiki_pages": wiki_pages_compiled,
        },
        "documents": doc_results,
    }


async def _parse_upload_file(f: UploadFile) -> str:
    """解析上传文件内容为纯文本。"""
    filename = (f.filename or "").lower()
    raw = await f.read()

    if filename.endswith((".txt", ".md")):
        # 尝试 utf-8，fallback gbk
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return raw.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return raw.decode("utf-8", errors="replace")

    if filename.endswith(".docx"):
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ValueError("需要安装 python-docx: pip install python-docx")

    if filename.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            import io
            pdf = fitz.open(stream=raw, filetype="pdf")
            text_parts = []
            for page in pdf:
                text_parts.append(page.get_text())
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError("需要安装 PyMuPDF: pip install pymupdf")

    # 其他格式尝试作为文本读取
    return raw.decode("utf-8", errors="replace")


# ── V14: 分析/确认 两阶段入库 ─────────────────────────


@router.post("/analyze")
async def analyze_files(
    files: list[UploadFile] = File(..., description="上传的文件列表"),
    project_id: str = Form(default="default", description="项目ID"),
) -> dict:
    """V14 Phase 1: 解析 + 蒸馏分析（不入库），返回分析结果供用户审核。

    流程: 上传 → 解析 → 蒸馏 → 返回结果 + batch_id
    用户审核后，调用 /finalize 按决策入库。
    """
    import hashlib
    import uuid
    from datetime import datetime, timezone
    from packages.distillation.pipeline import arun_pipeline
    from packages.common.types import RawDocument, SourceSystem
    from pathlib import Path

    domain_store = get_domain_store()
    project_domain_list = domain_store.get_refiner_domain_list(project_id=project_id)

    # Phase 1: 解析文件
    raw_docs: list[RawDocument] = []
    parse_errors: list[str] = []
    for f in files:
        try:
            content = await _parse_upload_file(f)
            if not content or len(content.strip()) < 10:
                parse_errors.append(f"{f.filename}: 内容为空或过短")
                continue
            content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
            name_hash = hashlib.md5((f.filename or '').encode()).hexdigest()[:8]
            doc_id = f"upload_{name_hash}_{content_hash}"
            raw_docs.append(RawDocument(
                doc_id=doc_id,
                title=Path(f.filename or "untitled").stem,
                content=content,
                source_system=SourceSystem.LOCAL if hasattr(SourceSystem, 'LOCAL') else SourceSystem.FEISHU,
                org_id=project_id,
                created_at=datetime.now(timezone.utc),
                updated_at=datetime.now(timezone.utc),
                file_size=len(content),
                metadata={"original_filename": f.filename},
            ))
        except Exception as e:
            parse_errors.append(f"{f.filename}: {str(e)}")

    if not raw_docs:
        raise HTTPException(status_code=400, detail=f"无可处理的文件。{'; '.join(parse_errors)}")

    # Phase 2: 保存到 RawStore（不可变层，分析阶段即保存）
    from api.deps import get_raw_store
    raw_store = get_raw_store()
    for doc in raw_docs:
        try:
            await raw_store.save_raw(doc, project_id=project_id)
        except Exception as e:
            log.warning("analyze_raw_store_save_failed", doc_id=doc.doc_id, error=str(e))

    # Phase 3: 蒸馏分析（不入库）
    batch = await arun_pipeline(raw_docs, domain_list_text=project_domain_list)

    # Phase 4: 构建分析结果
    doc_map = {d.doc_id: d for d in raw_docs}
    doc_results = []
    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        doc_results.append({
            "doc_id": r.doc_id,
            "title": doc.title if doc else r.doc_id,
            "decision": r.decision.value if r.decision else "UNKNOWN",
            "domain_id": r.refined_result.domain_id if r.refined_result else "",
            "summary": (r.refined_result.summary if r.refined_result else "")[:200],
            "doc_type": r.librarian_result.doc_type.value if r.librarian_result else "未知",
            "entity_count": len(r.refined_result.entities) if r.refined_result else 0,
            "keyword_count": len(r.refined_result.keywords) if r.refined_result else 0,
            "confidence": r.judge_result.confidence if r.judge_result else 0,
            "needs_review": r.needs_review,
        })

    # Phase 5: 暂存分析结果
    batch_id = str(uuid.uuid4())[:12]
    meta = get_metadata_store()
    meta.stage_batch(batch_id, {
        "batch": batch,
        "raw_docs": raw_docs,
        "project_id": project_id,
        "doc_map": doc_map,
    })

    log.info("analyze_completed", batch_id=batch_id, total=len(raw_docs),
             kept=batch.kept, archived=batch.archived)

    return {
        "status": "analyzed",
        "batch_id": batch_id,
        "total_uploaded": len(files),
        "parsed": len(raw_docs),
        "parse_errors": parse_errors,
        "distillation": {
            "kept": batch.kept,
            "archived": batch.archived,
            "discarded": batch.discarded,
        },
        "documents": doc_results,
    }


@router.post("/finalize")
async def finalize_batch(
    batch_id: str = Query(..., description="analyze 返回的 batch_id"),
    decisions: list[dict] = None,
) -> dict:
    """V14 Phase 2: 按用户审核决策执行入库。

    Body: [{"doc_id": "xxx", "decision": "KEEP"/"ARCHIVE"/"DISCARD"}, ...]
    未指定决策的文档使用 AI 自动判定结果。
    """
    from packages.common.types import Decision, DocumentCard
    from packages.storage.chunker import chunk_document
    from packages.storage.embedder import aembed_texts
    from packages.distillation.auto_cataloger import get_auto_cataloger

    meta = get_metadata_store()
    staged = meta.get_staged_batch(batch_id)
    if not staged:
        raise HTTPException(status_code=404, detail=f"分析批次 '{batch_id}' 不存在或已过期")

    batch = staged["batch"]
    raw_docs = staged["raw_docs"]
    project_id = staged["project_id"]
    doc_map = staged["doc_map"]

    # 解析用户决策覆盖
    user_decisions: dict[str, str] = {}
    if decisions:
        for d in decisions:
            if "doc_id" in d and "decision" in d:
                user_decisions[d["doc_id"]] = d["decision"]

    vec = get_vector_store()
    graph = get_graph_store()
    domain_store = get_domain_store()
    cataloger = get_auto_cataloger()

    total_chunks = 0
    kept_docs = 0
    archived_docs = 0
    all_bm25_entries: list[dict] = []

    for r in batch.results:
        doc = doc_map.get(r.doc_id)
        if not doc:
            continue

        # 用户决策优先，否则用 AI 决策
        final_decision = user_decisions.get(r.doc_id, r.decision.value if r.decision else "KEEP")

        if final_decision == "ARCHIVE" or final_decision == "DISCARD":
            archived_docs += 1
            continue

        # KEEP: 入库（复用 /ingest 的正确 API 签名）
        kept_docs += 1
        summary = r.refined_result.summary if r.refined_result else ""
        domain_id = r.refined_result.domain_id if r.refined_result else ""
        doc_type_str = r.librarian_result.doc_type.value if r.librarian_result else "未知"
        kw = ",".join(r.refined_result.keywords) if r.refined_result else ""

        cat = cataloger.generate_category_path(
            key_topics=r.librarian_result.key_topics if r.librarian_result else [],
            domain_id=domain_id,
            doc_type=doc_type_str,
        )

        # 存元数据（字段必须匹配 MetadataStore.upsert_document 的 "id" key）
        await meta.upsert_document({
            "id": doc.doc_id, "title": doc.title,
            "source_system": "local", "doc_type": doc_type_str,
            "version_id": None, "status": "ACTIVE",
            "decision": "KEEP",
            "kpi_retain": r.judge_result.kpi_retain if r.judge_result else None,
            "summary": summary, "keywords": kw,
            "category_path": cat, "org_id": project_id,
            "created_at": None, "updated_at": None,
            "access_level": "INTERNAL", "department_id": "",
            "judge_reasoning": r.judge_result.reasoning.model_dump() if r.judge_result and hasattr(r.judge_result, 'reasoning') and r.judge_result.reasoning else None,
        })

        # 向量化（使用正确的 chunk_document 签名）
        chunks = chunk_document(
            doc_id=doc.doc_id, content=doc.content,
            category_path=cat, doc_type=doc_type_str,
            source_system="local", updated_at=doc.updated_at,
            org_id=project_id, domain_id=domain_id,
        )
        if chunks:
            embeddings = await aembed_texts([c.content for c in chunks])
            for chunk, emb in zip(chunks, embeddings):
                chunk.embedding = emb
            await vec.insert_chunks(chunks)
            total_chunks += len(chunks)
            for c in chunks:
                all_bm25_entries.append({"chunk_id": c.chunk_id, "doc_id": c.doc_id, "content": c.content})

        # 图谱（使用 add_entities_and_relations 而非 add_node/add_edge）
        if r.refined_result:
            await graph.add_entities_and_relations(
                doc.doc_id, r.refined_result.entities, r.refined_result.relations,
                domain_id=domain_id,
            )
            doc_card = DocumentCard(
                doc_id=doc.doc_id, title=doc.title, domain_id=domain_id,
                description=r.refined_result.doc_description or summary[:200],
                key_elements=r.refined_result.key_elements,
                keywords=r.refined_result.keywords,
            )
            await domain_store.upsert_doc_card(doc_card, project_id=project_id)

    await graph.refresh_counts()

    # BM25 增量
    if all_bm25_entries:
        from api.deps import get_keyword_scorer
        get_keyword_scorer().add_chunks(all_bm25_entries)

    # Wiki 编译（使用过滤后的 results，尊重用户审核决策）
    finalized_results = [r for r in batch.results if user_decisions.get(r.doc_id, r.decision.value if r.decision else "KEEP") == "KEEP"]
    wiki_pages_compiled = 0
    try:
        wiki_pages_compiled = await _compile_wiki_for_batch(finalized_results, project_id, domain_store)
    except Exception as e:
        log.warning("finalize_wiki_compilation_failed", error=str(e))

    # 清理暂存
    meta.clear_staged_batch(batch_id)

    log.info("finalize_completed", batch_id=batch_id, kept=kept_docs,
             archived=archived_docs, chunks=total_chunks, wiki=wiki_pages_compiled)

    return {
        "status": "ok",
        "batch_id": batch_id,
        "kept": kept_docs,
        "archived": archived_docs,
        "vector_chunks": total_chunks,
        "wiki_pages": wiki_pages_compiled,
    }
