"""生成 3 个 .docx 格式的华能石化能源文档。"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parent.parent / "test_data" / "energy"


def _add_header(doc: Document, title: str, file_no: str, dept: str, author: str = ""):
    """添加公司标准文件头。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("华能石化有限公司")
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(16)

    info = doc.add_paragraph()
    info.add_run(f"文件编号：{file_no}\n").bold = False
    info.add_run(f"编制部门：{dept}\n")
    if author:
        info.add_run(f"编制人：{author}\n")
    info.add_run("审批人：总经理 李明远\n")
    info.add_run("发布日期：2025年02月20日\n")
    info.add_run("实施日期：2025年03月15日")


def create_doc16():
    """16_危化品仓储安全管理规定.docx — logistics > warehouse > hazmat_storage"""
    doc = Document()
    _add_header(doc, "危化品仓储安全管理规定", "HN-WL-016-2025",
                "物流管理部", "仓储管理科科长 王建华")

    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph(
        "第一条 为加强公司危险化学品仓储安全管理，预防仓储环节发生火灾、爆炸、泄漏、"
        "中毒等事故，根据《危险化学品安全管理条例》、GB 15603-2022《常用危险化学品储存通则》、"
        "GB 17914-2013《易燃易爆性商品储存养护技术条件》及《山东省危险化学品安全管理办法》，"
        "结合公司实际制定本规定。"
    )
    doc.add_paragraph(
        "第二条 本规定适用于公司原料库区、成品库区、中间罐区及化学试剂库的"
        "危险化学品储存管理。公司现有储罐区占地180亩，包括原油储罐8座（单罐容积"
        "10万立方米）、成品油储罐24座、液化气球罐4座、化工品储罐36座。"
    )

    doc.add_heading("第二章 储存分类", level=1)
    doc.add_paragraph(
        "第三条 危险化学品按储存方式分为三类："
    )
    # 添加分类表格
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["储存类别", "适用品类", "储存条件", "安全间距"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["甲类仓库", "液化石油气、汽油、石脑油、苯", "防爆电气、机械通风、可燃气体报警",
         "距办公区≥50m"],
        ["乙类仓库", "柴油、润滑油、沥青", "防火分区、自动喷淋灭火", "距办公区≥30m"],
        ["丙类仓库", "烧碱、硫酸（稀）、助剂", "防腐蚀地面、应急冲洗设施", "距办公区≥25m"],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_heading("第三章 禁忌物料管理", level=1)
    doc.add_paragraph(
        "第四条 性质互相矛盾的危险化学品严禁同库储存。公司主要禁忌组合如下："
    )
    table2 = doc.add_table(rows=5, cols=3, style="Table Grid")
    h2 = ["物料A", "物料B", "禁忌原因"]
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
    pairs = [
        ["汽油（易燃液体）", "双氧水（氧化剂）", "氧化剂加速易燃液体燃烧"],
        ["硫酸（强酸）", "烧碱溶液（强碱）", "酸碱反应剧烈放热"],
        ["液化石油气（压缩气体）", "硫酸（腐蚀品）", "泄漏后腐蚀容器加剧风险"],
        ["苯（有毒易燃）", "盐酸（腐蚀品）", "混合产生有毒气体"],
    ]
    for r, row_data in enumerate(pairs):
        for c, val in enumerate(row_data):
            table2.rows[r + 1].cells[c].text = val

    doc.add_heading("第四章 温湿度控制", level=1)
    doc.add_paragraph(
        "第五条 各类仓库温湿度控制标准：\n"
        "（一）甲类仓库：温度≤30℃，相对湿度≤75%，配备温湿度在线监测和自动报警\n"
        "（二）乙类仓库：温度≤35℃，相对湿度≤80%\n"
        "（三）液化气球罐区：罐壁温度≤45℃，超温启动喷淋降温\n"
        "（四）化学试剂库：温度15-25℃，相对湿度45-65%，精密空调控温"
    )
    doc.add_paragraph(
        "第六条 夏季高温期间（6-9月）每2小时巡检一次库区温度，"
        "温度超过预警值时启动降温措施（通风、喷淋、遮阳）。"
    )

    doc.add_heading("第五章 库区安全设施", level=1)
    doc.add_paragraph(
        "第七条 危化品库区必须配备以下安全设施：\n"
        "（一）可燃气体/有毒气体检测报警装置，报警值设置：LEL 25%（预警）/50%（报警）\n"
        "（二）自动消防系统（泡沫灭火系统+干粉灭火系统）\n"
        "（三）防雷接地系统，接地电阻≤10Ω\n"
        "（四）防静电设施，人体静电消除器设在入口处\n"
        "（五）应急照明（不低于30分钟）和疏散指示标志\n"
        "（六）围堰/防火堤，容积不小于最大单罐容积的100%\n"
        "（七）视频监控系统，覆盖率100%，录像保存≥90天"
    )

    doc.add_heading("第六章 出入库管理", level=1)
    doc.add_paragraph(
        "第八条 危化品出入库执行双人验收制度：\n"
        "（一）入库时核对品名、规格、数量、包装完好性、安全数据表（SDS）\n"
        "（二）按品类分区、分垛存放，堆垛高度不超过2米\n"
        "（三）储存期限不超过6个月，先进先出\n"
        "（四）出库时再次核对，记录去向和用途\n"
        "（五）每月末盘点，账物相符率≥99.5%"
    )

    doc.add_heading("第七章 附则", level=1)
    doc.add_paragraph("第九条 本规定由物流管理部负责解释和修订。")
    doc.add_paragraph("第十条 本规定自发布之日起施行。")

    path = OUT_DIR / "16_危化品仓储安全管理规定.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


def create_doc17():
    """17_供应商准入与绩效评价管理办法.docx — procurement > supplier"""
    doc = Document()
    _add_header(doc, "供应商准入与绩效评价管理办法", "HN-CG-017-2025",
                "采购管理部", "采购管理部部长 赵国栋")

    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph(
        "第一条 为规范公司供应商管理，建立公平、公正、透明的供应商准入与退出机制，"
        "持续优化供应链质量，根据《中华人民共和国招标投标法》、"
        "中国石油化工集团公司《物资供应商管理办法》及公司采购管理制度，制定本办法。"
    )
    doc.add_paragraph(
        "第二条 本办法适用于向公司供应设备、材料、备件、化工原料、安全防护用品、"
        "劳保用品及技术服务的所有供应商。"
    )
    doc.add_paragraph(
        "第三条 采购管理部是供应商管理的归口部门，各使用部门配合提供技术评审意见。"
    )

    doc.add_heading("第二章 供应商准入", level=1)
    doc.add_paragraph(
        "第四条 新供应商准入须提交以下资质材料：\n"
        "（一）营业执照（经营范围覆盖供货品类）\n"
        "（二）税务登记证或统一社会信用代码证\n"
        "（三）质量管理体系认证证书（ISO 9001或等效）\n"
        "（四）安全生产许可证（适用时）\n"
        "（五）特种设备制造/安装许可证（适用时）\n"
        "（六）近三年业绩证明（同类项目合同或中标通知书）\n"
        "（七）财务报表（最近一个年度）\n"
        "（八）无重大违法违规声明"
    )
    doc.add_paragraph(
        "第五条 供应商资质审查由采购管理部组织，技术部门、安全部门、"
        "质量部门参与联合评审，形成《供应商资质评审报告》。"
    )

    doc.add_heading("第三章 绩效评分模型", level=1)
    doc.add_paragraph("第六条 供应商年度绩效评价采用百分制，评分维度及权重如下：")

    table = doc.add_table(rows=7, cols=4, style="Table Grid")
    headers = ["评分维度", "权重", "评分标准", "数据来源"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ["产品质量", "35%", "合格率≥99%得满分，每降0.5%扣5分", "质检报告"],
        ["交货及时性", "25%", "准时交货率≥95%得满分，每降2%扣5分", "入库记录"],
        ["价格竞争力", "15%", "低于市场均价得满分，高于均价按比例扣分", "市场调研"],
        ["售后服务", "10%", "响应时间≤4h满分，超时按比例扣分", "服务工单"],
        ["安全合规", "10%", "零安全事故满分，发生事故视严重程度扣分", "安全记录"],
        ["创新能力", "5%", "提供技术改进方案或新产品推荐", "技术部门评价"],
    ]
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_heading("第四章 供应商分级", level=1)
    doc.add_paragraph("第七条 根据年度绩效评分，供应商分为四个等级：")

    table2 = doc.add_table(rows=5, cols=4, style="Table Grid")
    h2 = ["等级", "评分区间", "管理策略", "采购份额"]
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
    levels = [
        ["A级（战略）", "90-100分", "长期合作、优先采购、年度框架协议", "≥40%"],
        ["B级（优选）", "75-89分", "正常合作、竞争性谈判", "30-40%"],
        ["C级（合格）", "60-74分", "限制合作、整改观察、减少份额", "10-20%"],
        ["D级（淘汰）", "<60分", "暂停合作、列入黑名单", "0%"],
    ]
    for r, row_data in enumerate(levels):
        for c, val in enumerate(row_data):
            table2.rows[r + 1].cells[c].text = val

    doc.add_paragraph(
        "第八条 D级供应商立即暂停供货资格，列入供应商黑名单，"
        "2年内不得重新申请准入。C级供应商给予6个月整改期，"
        "整改后重新评审，仍达不到B级标准的降为D级。"
    )

    doc.add_heading("第五章 合同管理", level=1)
    doc.add_paragraph(
        "第九条 采购合同签订须遵循以下原则：\n"
        "（一）A级供应商可签订年度框架合同，有效期1-3年\n"
        "（二）单次采购金额≥50万元须经招标或竞争性谈判\n"
        "（三）合同必须包含质量保证条款、违约责任和争议解决机制\n"
        "（四）涉及危化品的合同须附《安全技术协议》\n"
        "（五）合同原件由采购管理部统一归档，扫描件录入ERP系统"
    )

    doc.add_heading("第六章 附则", level=1)
    doc.add_paragraph("第十条 本办法由采购管理部负责解释和修订。")
    doc.add_paragraph("第十一条 本办法自发布之日起施行。")

    path = OUT_DIR / "17_供应商准入与绩效评价管理办法.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


def create_doc18():
    """18_应急演练计划与评估管理办法.docx — emergency > drill"""
    doc = Document()
    _add_header(doc, "应急演练计划与评估管理办法", "HN-YJ-018-2025",
                "应急管理部", "应急管理部部长 孙伟明")

    doc.add_heading("第一章 总则", level=1)
    doc.add_paragraph(
        "第一条 为提高公司突发事件应急处置能力，检验应急预案的实用性和有效性，"
        "根据《生产安全事故应急条例》（国务院令第708号）、AQ/T 9007-2019"
        "《生产安全事故应急演练基本规范》、GB/T 29639-2020"
        "《生产经营单位生产安全事故应急预案编制导则》，制定本办法。"
    )
    doc.add_paragraph(
        "第二条 本办法适用于公司及各下属单位组织的综合应急演练、专项应急演练"
        "和现场处置演练。"
    )

    doc.add_heading("第二章 年度演练计划", level=1)
    doc.add_paragraph(
        "第三条 应急管理部于每年12月编制下年度《应急演练计划》，经分管副总审批后下发。"
    )
    doc.add_paragraph("第四条 年度演练频次要求：")

    table = doc.add_table(rows=5, cols=5, style="Table Grid")
    headers = ["演练类型", "频次", "参演范围", "持续时间", "评审要求"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ["综合应急演练", "≥1次/年", "全公司", "2-4小时", "外部专家评审"],
        ["专项应急演练", "≥2次/年", "相关部门+消防", "1-2小时", "公司级评审"],
        ["现场处置演练", "≥4次/年", "班组/车间", "30-60分钟", "部门级评审"],
        ["桌面推演", "≥2次/年", "应急指挥部", "1-2小时", "内部评审"],
    ]
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    doc.add_heading("第三章 演练情景设计", level=1)
    doc.add_paragraph(
        "第五条 演练情景应覆盖公司主要风险类型，2025年度重点演练情景包括："
    )
    scenarios = [
        "催化裂化装置油气泄漏引发火灾爆炸（综合演练，6月）",
        "液化石油气球罐泄漏及毒气扩散（专项演练，4月）",
        "储罐区全面火灾（专项演练，9月）",
        "危化品运输车辆翻覆泄漏（现场演练，3月、7月）",
        "废水处理场超标排放应急（现场演练，5月、10月）",
        "地震导致装置损坏（桌面推演，8月）",
        "暴雨洪涝导致罐区进水（桌面推演，11月）",
    ]
    for s in scenarios:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph(
        "第六条 演练情景设计须包含以下要素：\n"
        "（一）事故类型及危害程度假设\n"
        "（二）事故发生的时间、地点、气象条件\n"
        "（三）事故影响范围和可能的后果\n"
        "（四）需要调动的应急资源\n"
        "（五）模拟伤亡情况（如有）"
    )

    doc.add_heading("第四章 演练组织实施", level=1)
    doc.add_paragraph(
        "第七条 演练组织机构：\n"
        "（一）总指挥：分管安全副总经理\n"
        "（二）副总指挥：安全生产部部长、应急管理部部长\n"
        "（三）演练策划组：编制演练方案、协调资源\n"
        "（四）演练执行组：现场导调、信号控制\n"
        "（五）演练评估组：记录观察、拍照录像\n"
        "（六）后勤保障组：医疗、通讯、物资保障"
    )
    doc.add_paragraph(
        "第八条 演练实施流程：\n"
        "（一）演练前7天发布演练通知\n"
        "（二）演练前3天组织桌面推演和关键岗位培训\n"
        "（三）演练前1天检查应急物资和设备就位\n"
        "（四）演练当天按方案执行，评估组全程记录\n"
        "（五）演练结束后立即召开现场总结会"
    )

    doc.add_heading("第五章 演练评估标准", level=1)
    doc.add_paragraph("第九条 演练效果评估采用百分制：")

    table2 = doc.add_table(rows=7, cols=3, style="Table Grid")
    h2 = ["评估项目", "分值", "评分标准"]
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
    items = [
        ["报警与通讯", "15分", "报警及时（≤3分钟）、通讯畅通、信息准确"],
        ["应急响应", "20分", "响应时间≤5分钟、指挥到位、职责清晰"],
        ["现场处置", "25分", "处置措施正确、操作规范、无二次伤害"],
        ["人员疏散", "15分", "疏散有序、清点到位、用时≤10分钟"],
        ["医疗救护", "10分", "救护及时、处置正确"],
        ["后勤保障", "15分", "物资到位、协调有序、记录完整"],
    ]
    for r, row_data in enumerate(items):
        for c, val in enumerate(row_data):
            table2.rows[r + 1].cells[c].text = val

    doc.add_paragraph(
        "第十条 评估结果分级：\n"
        "≥90分：优秀——应急预案有效，可操作性强\n"
        "75-89分：良好——基本有效，局部需改进\n"
        "60-74分：合格——存在明显不足，须在30天内整改\n"
        "<60分：不合格——须在15天内重新组织演练"
    )

    doc.add_heading("第六章 改进跟踪", level=1)
    doc.add_paragraph(
        "第十一条 演练结束后10个工作日内，应急管理部编制《演练评估报告》，内容包括：\n"
        "（一）演练基本情况（时间、地点、参演人数）\n"
        "（二）各环节评分及扣分原因\n"
        "（三）发现的问题清单\n"
        "（四）整改措施及责任人\n"
        "（五）预案修订建议"
    )
    doc.add_paragraph(
        "第十二条 问题整改采用PDCA闭环管理：\n"
        "（一）Plan：制定整改方案，明确完成时限\n"
        "（二）Do：责任部门落实整改\n"
        "（三）Check：应急管理部验收整改效果\n"
        "（四）Act：将改进措施纳入应急预案修订"
    )

    doc.add_heading("第七章 附则", level=1)
    doc.add_paragraph("第十三条 本办法由应急管理部负责解释和修订。")
    doc.add_paragraph("第十四条 本办法自发布之日起施行。")

    path = OUT_DIR / "18_应急演练计划与评估管理办法.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


if __name__ == "__main__":
    print("Generating .docx energy documents ...")
    create_doc16()
    create_doc17()
    create_doc18()
    print("Done! 3 docx files created in test_data/energy/")
